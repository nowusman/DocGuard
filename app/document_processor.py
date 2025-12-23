# document_processor.py

import base64
import hashlib
import io
import json
import os
import re
import html
from collections import OrderedDict
from concurrent.futures import ThreadPoolExecutor, as_completed
from copy import deepcopy
from datetime import datetime

import fitz  # PyMuPDF
import numpy as np
import spacy
import zipfile
import xml.etree.ElementTree as ET
from PIL import Image as PILImage, ImageStat
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from time import perf_counter

from config import (
    PII_PATTERNS,
    JSON_SCHEMA,
    OCR_CONFIG,
    PDF_HEADER_RATIO,
    THROUGHPUT_MODE,
    OCR_MAX_IMAGES_PER_DOC,
    OCR_RENDER_SCALE,
    VERBOSE_LOGGING,
    PDF_ENGINE,
    OCR_ENGINE,
    MAX_CACHE_ITEMS,
)

class DocumentProcessor:
    def __init__(self):
        self.verbose_logging = VERBOSE_LOGGING
        self.throughput_mode = THROUGHPUT_MODE
        self.ocr_enabled = bool(OCR_CONFIG.get('enabled', True))
        self._timing = {}
        self._ocr_images_processed = 0
        self._ocr_images_skipped = 0
        self.max_cache_items = MAX_CACHE_ITEMS
        self._cache = OrderedDict()
        self._paddle_ocr = None
        self.ocr_engine = OCR_ENGINE
        self.anonymize_terms = []
        self.anonymize_replace = ""
        self._anonymize_terms_regex = None
        # Load spaCy model
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            print("Warning: spaCy model 'en_core_web_sm' not found. Please install it.")
            self.nlp = None
        
        # Check OCR availability
        self.ocr_available = self._check_ocr_availability()

    def _log(self, message: str):
        if self.verbose_logging:
            print(message)

    def _reset_timing(self):
        self._timing = {}
        self._ocr_images_processed = 0
        self._ocr_images_skipped = 0

    def _record_timing(self, key: str, duration: float):
        if duration is None:
            return
        self._timing[key] = self._timing.get(key, 0.0) + float(duration)

    def _build_cache_key(self, file_content, anonymize, remove_pii, extract_json, options):
        if not self.max_cache_items:
            return None
        hasher = hashlib.sha256()
        if isinstance(file_content, (bytes, bytearray)):
            file_bytes = bytes(file_content)
        else:
            file_bytes = str(file_content).encode('utf-8', errors='ignore')
        hasher.update(file_bytes)
        normalized = {
            'anonymize': bool(anonymize),
            'remove_pii': bool(remove_pii),
            'extract_json': bool(extract_json),
            'throughput_mode': bool(self.throughput_mode),
            'ocr_enabled': bool(self.ocr_enabled),
            'options': options or {}
        }
        hasher.update(json.dumps(normalized, sort_keys=True, default=str).encode('utf-8'))
        return hasher.hexdigest()

    def _get_cached_result(self, cache_key):
        if not cache_key or not self._cache:
            return None
        cached = self._cache.get(cache_key)
        if cached:
            self._cache.move_to_end(cache_key)
        return cached

    def _store_cache_result(self, cache_key, result):
        if not cache_key or not self.max_cache_items:
            return
        content, extension, metadata = result
        self._cache[cache_key] = (content, extension, deepcopy(metadata))
        self._cache.move_to_end(cache_key)
        if len(self._cache) > self.max_cache_items:
            self._cache.popitem(last=False)

    def _finalize_metadata(self, metadata):
        metadata = metadata or {}
        metadata['timing'] = dict(self._timing)
        metadata['throughput_mode'] = self.throughput_mode
        metadata['cache_hit'] = metadata.get('cache_hit', False)
        metadata['ner_mode'] = 'regex_only' if (self.throughput_mode or not self.nlp) else 'spacy_batch'
        metadata['ocr'] = {
            'engine': self.ocr_engine if self.ocr_available else 'unavailable',
            'images_processed': self._ocr_images_processed,
            'images_skipped': self._ocr_images_skipped,
            'max_images_per_doc': OCR_MAX_IMAGES_PER_DOC,
            'enabled': bool(self.ocr_enabled) and not self.throughput_mode,
        }
        return metadata

    def _finalize_result(self, content, extension, metadata):
        return content, extension, self._finalize_metadata(metadata)

    def _finalize_with_cache(self, content, extension, metadata, cache_key):
        result = self._finalize_result(content, extension, metadata)
        if cache_key:
            self._store_cache_result(cache_key, result)
        return result
    
    def _check_ocr_availability(self):
        """
        Check OCR availability
        """
        try:
            from paddleocr import PaddleOCR

            self._paddle_ocr = PaddleOCR(
                lang="en",
                use_angle_cls=False,
                show_log=False,
                use_gpu=False,
            )
            self.ocr_engine = "paddle"
            return True
        except ImportError:
            print("Warning: PaddleOCR not available. Install paddleocr for OCR support.")
            self._paddle_ocr = None
            return False
        except Exception as exc:
            print(f"Warning: PaddleOCR initialization failed: {exc}")
            self._paddle_ocr = None
            return False
    
    def _has_table_indicators(self, text: str) -> bool:
        """
        Perform fast heuristic check for table indicators in text.
        Returns True if text contains patterns that suggest tables might be present.
        """
        if not text:
            return False
        
        # Common table indicators
        table_indicators = [
            # Grid-like patterns
            r'\|\s*[\w\s]+\s*\|',  # Pipe separators
            r'\+[-]+\+',  # ASCII table borders
            r'[\w\s]+\s+\|\s+[\w\s]+',  # Text with pipe separator
            r'\b(table|tab\.?|tbl)\b',  # Table references
            # Column-like patterns
            r'\s{4,}[\w\s]+\s{4,}[\w\s]+',  # Multiple spaces as column separators
            r'\t+[\w\s]+\t+[\w\s]+',  # Tabs as column separators
        ]
        
        # Check each pattern
        for pattern in table_indicators:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        
        # Check for numeric data in rows (potential data tables)
        lines = text.split('\n')
        table_like_lines = 0
        
        for line in lines[:50]:  # Check first 50 lines only for speed
            line = line.strip()
            if not line:
                continue
            
            # Check if line looks like table row
            # Multiple items separated by consistent delimiters
            if '|' in line and line.count('|') >= 2:
                table_like_lines += 1
            elif len(re.split(r'\s{3,}', line)) >= 3:  # Multiple columns with wide spacing
                table_like_lines += 1
            elif re.search(r'\d+[,\d]*\s+\d+[,\d]*', line):  # Multiple numbers
                table_like_lines += 1
        
        # If we found multiple table-like lines, likely has tables
        return table_like_lines >= 3
    
    def process_document(self, file_content, filename, anonymize=False, remove_pii=False, extract_json=False, options=None):
        """
        The main functions for processing documents
        """
        options = dict(options) if options else {}
        self.verbose_logging = options.get('verbose_logging', VERBOSE_LOGGING)
        self.throughput_mode = options.get('throughput_mode', THROUGHPUT_MODE)
        if 'ocr_enabled' in options:
            self.ocr_enabled = bool(options['ocr_enabled'])
        self._set_anonymization_settings(
            options.get('anonymize_terms'),
            options.get('anonymize_replace'),
        )
        options['anonymize_terms'] = self.anonymize_terms
        options['anonymize_replace'] = self.anonymize_replace
        self._reset_timing()
        cache_key = None
        cached_result = None
        if self.max_cache_items:
            cache_key = self._build_cache_key(
                file_content,
                anonymize,
                remove_pii,
                extract_json,
                options,
            )
            cached_result = self._get_cached_result(cache_key)
            if cached_result:
                cached_content, cached_extension, cached_metadata = cached_result
                metadata_copy = deepcopy(cached_metadata)
                metadata_copy['cache_hit'] = True
                return cached_content, cached_extension, metadata_copy

        file_extension = f".{filename.split('.')[-1].lower()}"
        
        # Read document content
        if file_extension == '.txt':
            content, metadata = self._read_txt(file_content)
        elif file_extension == '.docx':
            if anonymize or remove_pii:
                content, metadata = None, None
            else:
                content, metadata = self._read_docx(file_content)
        elif file_extension == '.pdf':
            content, metadata = self._read_pdf(file_content)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Processing Options
        processing_info = {
            'anonymized': anonymize,
            'pii_removed': remove_pii,
            'extracted_to_json': extract_json
        }
        
        original_content = content
        
        # DOCX files, handle byte data when applying processing options
        if file_extension == '.docx':
            if anonymize or remove_pii:
                processed_bytes = file_content

                if anonymize:
                    processed_bytes = self._process_docx_xml(processed_bytes, 'anonymize')

                if remove_pii:
                    processed_bytes = self._process_docx_xml(processed_bytes, 'remove_pii')

                doc = Document(io.BytesIO(processed_bytes))
                content, tables_data, paragraphs = self._extract_docx_text_structures(doc)
                images_data = self._extract_images_from_docx(doc, processed_bytes)
                metadata = {
                    'text_content': content,
                    'tables': tables_data,
                    'images': images_data,
                    'paragraphs': paragraphs,
                }

                if extract_json:
                    result = self._extract_to_json(content, filename, file_extension, processing_info, metadata, file_content)
                    return self._finalize_with_cache(result, '.json', metadata, cache_key)
                else:
                    pdf_content = self._create_pdf_with_layout(content, filename, file_content, metadata)
                    return self._finalize_with_cache(pdf_content, '.pdf', metadata, cache_key)

            if extract_json:
                result = self._extract_to_json(content, filename, file_extension, processing_info, metadata, file_content)
                return self._finalize_with_cache(result, '.json', metadata, cache_key)
            else:
                pdf_content = self._create_pdf_with_layout(content, filename, file_content, metadata)
                return self._finalize_with_cache(pdf_content, '.pdf', metadata, cache_key)
        
        else:
            # TXT and PDF files, use the existing string processing logic
            if anonymize:
                content = self._apply_anonymization(content, file_extension, file_content)
            
            if remove_pii:
                content = self._remove_pii(content, file_extension, file_content)
            
            if extract_json:
                result = self._extract_to_json(content, filename, file_extension, processing_info, metadata, file_content)
                return self._finalize_with_cache(result, '.json', metadata, cache_key)
            else:
                if anonymize or remove_pii:
                    if file_extension == '.txt':
                        pdf_content = self._create_pdf(content, filename)
                        return self._finalize_with_cache(pdf_content, '.pdf', metadata, cache_key)
                    else:
                        pdf_content = self._create_pdf_with_layout(content, filename, file_content, metadata)
                        return self._finalize_with_cache(pdf_content, '.pdf', metadata, cache_key)
                else:
                    return self._finalize_with_cache(file_content, file_extension, metadata, cache_key)
    
    def _read_txt(self, file_content):
        """Read TXT file"""
        start_time = perf_counter()
        if isinstance(file_content, bytes):
            content = file_content.decode('utf-8')
        else:
            content = file_content
            
        metadata = {
            'text_content': content,
            'tables': [],
            'images': []
        }
        self._record_timing('read_txt', perf_counter() - start_time)
        return content, metadata
    
    def _read_docx(self, file_content):
        """Read DOCX file"""
        start_time = perf_counter()
        doc = Document(io.BytesIO(file_content))
        content, tables_data, paragraphs = self._extract_docx_text_structures(doc)
        images_data = self._extract_images_from_docx(doc, file_content)
        metadata = {
            'text_content': content,
            'tables': tables_data,
            'images': images_data,
            'paragraphs': paragraphs,
        }
        
        self._record_timing('read_docx', perf_counter() - start_time)
        return content, metadata

    def _extract_docx_text_structures(self, doc):
        """Extract docx text, tables, and paragraphs in a reusable way."""
        full_text = []
        tables_data = []
        paragraphs = []

        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)

        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables_data.append({
                'table_index': table_idx,
                'data': table_data,
                'rows': len(table.rows),
                'cols': len(table.columns) if hasattr(table, 'columns') else len(table.rows[0].cells) if table.rows else 0
            })

            for row in table_data:
                full_text.append(' | '.join(row))

        content = '\n'.join(full_text)
        return content, tables_data, paragraphs
    

    #############################3
    def _read_pdf(self, file_content):
        """Entry point for PDF reads with PyMuPDF single-pass."""
        try:
            return self._read_pdf_optimized(file_content)
        except Exception as exc:
            self._log(f"PyMuPDF single-pass failed, returning text-only fallback: {exc}")
            return self._read_pdf_text_only(file_content)

    def _read_pdf_optimized(self, file_content):
        """Single-pass PDF extraction using PyMuPDF only."""
        read_start = perf_counter()
        text_chunks = []
        tables_data = []
        images_data = []
        table_index = 0
        header_ratio = PDF_HEADER_RATIO or 0
        file_bytes = bytes(file_content) if isinstance(file_content, (bytes, bytearray)) else str(file_content).encode('utf-8', errors='ignore')

        with fitz.open(stream=file_bytes, filetype="pdf") as pdf_doc:
            for page_num, page in enumerate(pdf_doc):
                rect = page.rect
                clip = rect
                if rect is not None and rect.height > 0:
                    top_offset = rect.height * header_ratio
                    bottom_offset = rect.height * (1 - header_ratio)
                    if bottom_offset > top_offset:
                        clip = fitz.Rect(rect.x0, rect.y0 + top_offset, rect.x1, rect.y0 + bottom_offset)
                try:
                    page_text = page.get_text("text", clip=clip)
                except TypeError:
                    page_text = page.get_text(clip=clip)
                if page_text:
                    text_chunks.append(page_text)

                # Check if we should attempt table extraction
                should_extract_tables = not self.throughput_mode
                
                if should_extract_tables:
                    # Fast heuristic check: skip expensive table extraction if no table indicators found
                    has_table_indicators = self._has_table_indicators(page_text)
                    
                    if not has_table_indicators:
                        self._log(f"Page {page_num + 1}: No table indicators found, skipping table extraction")
                    else:
                        self._log(f"Page {page_num + 1}: Table indicators found, proceeding with extraction")
                        
                        tables_start = perf_counter()
                        # Use PyMuPDF for table extraction
                        new_tables, table_index = self._extract_tables_with_pymupdf(page, page_num, table_index)
                        if new_tables:
                            tables_data.extend(new_tables)
                            self._timing['table_extraction'] = self._timing.get('table_extraction', 0) + (perf_counter() - tables_start)

                page_images = self._extract_images_with_pymupdf(pdf_doc, page, page_num)
                if page_images:
                    images_data.extend(page_images)

        text = "\n".join(text_chunks)
        images_data = self._process_images_with_ocr(images_data)
        metadata = {
            'text_content': text,
            'tables': tables_data,
            'images': images_data,
            'pdf_engine': f"{PDF_ENGINE}_single_pass"
        }
        self._record_timing('read_pdf', perf_counter() - read_start)
        return text, metadata
    
    def _read_pdf_text_only(self, file_content):
        """Lightweight text-only fallback using PyMuPDF without tables/images."""
        start_time = perf_counter()
        text_chunks = []
        file_bytes = bytes(file_content) if isinstance(file_content, (bytes, bytearray)) else str(file_content).encode('utf-8', errors='ignore')
        try:
            with fitz.open(stream=file_bytes, filetype="pdf") as pdf_doc:
                for page in pdf_doc:
                    try:
                        page_text = page.get_text("text")
                    except TypeError:
                        page_text = page.get_text()
                    if page_text:
                        text_chunks.append(page_text)
        except Exception as exc:
            raise ValueError(f"Unable to read PDF with PyMuPDF: {exc}") from exc

        text = "\n".join(text_chunks)
        metadata = {
            'text_content': text,
            'tables': [],
            'images': [],
            'pdf_engine': f"{PDF_ENGINE}_text_only"
        }
        self._record_timing('read_pdf', perf_counter() - start_time)
        return text, metadata

    def _extract_tables_with_pymupdf(self, page, page_num, start_index):
        tables = []
        table_offset = 0
        try:
            finder = page.find_tables()
            if finder is None:
                return tables, start_index
            if hasattr(finder, 'tables'):
                table_list = finder.tables or []
            elif isinstance(finder, (list, tuple)):
                table_list = list(finder)
            else:
                table_list = [finder]
            for table in table_list:
                try:
                    data = table.extract()
                except Exception:
                    data = table.to_list() if hasattr(table, 'to_list') else None
                if not data:
                    continue
                rows = len(data)
                cols = len(data[0]) if rows else 0
                tables.append({
                    'table_index': start_index + table_offset,
                    'data': data,
                    'rows': rows,
                    'cols': cols,
                    'page': page_num + 1,
                    'extraction_method': 'pymupdf'
                })
                table_offset += 1
        except Exception as exc:
            self._log(f"PyMuPDF table extraction failed on page {page_num + 1}: {exc}")
        return tables, start_index + table_offset

    def _extract_images_with_pymupdf(self, doc, page, page_num):
        images = []
        seen_xrefs = set()
        try:
            image_list = page.get_images(full=True) or []
        except Exception as exc:
            self._log(f"PyMuPDF image listing failed on page {page_num + 1}: {exc}")
            return images
        for img_idx, img in enumerate(image_list):
            xref = img[0]
            if xref in seen_xrefs:
                continue
            seen_xrefs.add(xref)
            pix = self._extract_image_with_pymupdf(doc, page_num, img_idx, bbox=None, xref=xref)
            if not pix:
                continue
            try:
                image_bytes = pix.tobytes("png")
            except Exception:
                continue
            finally:
                pix = None
            image_info = {
                "page": page_num + 1,
                "type": "pdf_embedded_image",
                "description": f"Image on page {page_num + 1}",
                "image_data": image_bytes,
                "image_format": self._get_image_format(image_bytes),
                "extracted_text": "",
                "ocr_applied": False,
            }
            images.append(image_info)
        return images

    
    def _normalize_anonymization_terms(self, terms) -> list:
        """Normalize anonymization terms: trim, drop empties, de-duplicate preserving order."""
        if not terms:
            return []
        normalized = []
        seen = set()
        for raw_term in terms:
            if raw_term is None:
                continue
            term = str(raw_term).strip()
            if not term:
                continue
            key = term.lower()
            if key in seen:
                continue
            seen.add(key)
            normalized.append(term)
        return normalized

    def _set_anonymization_settings(self, terms_override, replace_override):
        """Apply per-request anonymization overrides with config defaults."""
        terms_source = terms_override if terms_override is not None else []
        replacement_source = replace_override if replace_override is not None else ""
        self.anonymize_terms = self._normalize_anonymization_terms(terms_source)
        self.anonymize_replace = "" if replacement_source is None else str(replacement_source)
        if self.anonymize_terms:
            pattern = "|".join(map(re.escape, self.anonymize_terms))
            self._anonymize_terms_regex = re.compile(pattern, re.IGNORECASE)
        else:
            self._anonymize_terms_regex = None

    def _apply_anonymization_terms(self, text: str) -> str:
        """Apply configured anonymization terms to a text snippet."""
        if not self._anonymize_terms_regex:
            return text
        replacement_value = " " if self.anonymize_replace == "" else self.anonymize_replace
        return self._anonymize_terms_regex.sub(replacement_value, text)

    def _apply_anonymization(self, content, file_extension, original_content):
        """Application anonymization processing - only for TXT and PDF"""
        
        if isinstance(content, bytes):
            content = content.decode('utf-8')
            
        return self._apply_anonymization_terms(content)
    
    def _remove_pii_fast(self, content: str) -> str:
        """Regex-only PII removal for fast mode or as a preprocessing step."""
        pii_removed_content = content
        for pattern in PII_PATTERNS.values():
            pii_removed_content = pattern.sub('[PII_REMOVED]', pii_removed_content)
        return pii_removed_content
    
    def _apply_spacy_entities_batch(self, texts, regex_cleaned=False):
        """Apply spaCy entity redaction in batch for PERSON/ORG/GPE."""
        if not texts:
            return []
        if not self.nlp:
            if regex_cleaned:
                return list(texts)
            return [self._remove_pii_fast(text) for text in texts]

        redacted = []
        for doc, original in zip(self.nlp.pipe(texts, batch_size=50, n_process=1), texts):
            base = original if regex_cleaned else self._remove_pii_fast(original)
            redacted.append(self._apply_ner_spans(base, doc))
        return redacted

    def _apply_ner_spans(self, text, doc):
        """Redact spaCy entity spans without over-replacing substrings."""
        if not getattr(doc, "ents", None):
            return text
        spans = [
            (ent.start_char, ent.end_char)
            for ent in doc.ents
            if ent.label_ in {"PERSON", "ORG", "GPE"}
        ]
        if not spans:
            return text
        spans.sort()
        out = []
        last = 0
        for start, end in spans:
            out.append(text[last:start])
            out.append("[PII_REMOVED]")
            last = end
        out.append(text[last:])
        return "".join(out)

    def _process_text_batch(self, texts, operation):
        """Process a list of text snippets in batch using regex + spaCy pipe."""
        processed = []
        if operation == 'anonymize':
            return [self._apply_anonymization_terms(text) for text in texts]

        if operation == 'remove_pii':
            regex_only = [self._remove_pii_fast(text) for text in texts]
            if self.throughput_mode or not self.nlp:
                return regex_only
            return self._apply_spacy_entities_batch(regex_only, regex_cleaned=True)

        return texts

    def _apply_text_to_paragraph(self, paragraph, new_text: str):
        """Replace paragraph content with processed text while keeping styling simple."""
        for run in paragraph.runs:
            run.text = ""
        
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)
    
    def _remove_pii(self, content, file_extension, original_content):
        """Remove PII information - for TXT and PDF only"""
        start_time = perf_counter()
        
        if isinstance(content, bytes):
            content = content.decode('utf-8')
            
        regex_cleaned = self._remove_pii_fast(content)
        
        if not self.throughput_mode and self.nlp:
            regex_cleaned = self._apply_spacy_entities_batch([regex_cleaned], regex_cleaned=True)[0]
        
        self._record_timing('pii_removal', perf_counter() - start_time)
        return regex_cleaned
    
    def _process_docx_xml(self, file_content, operation):
        """
        Processing XML Content of DOCX Files
        """
        try:
            return self._process_docx_direct(file_content, operation)
        except Exception as e:
            self._log(f"Error in direct DOCX processing: {e}, falling back to python-docx")
            return self._process_docx_with_python_docx(file_content, operation)
    
    def _process_docx_direct(self, file_content, operation):
        """
        Process DOCX
        """
        try:
            input_zip = io.BytesIO(file_content)
            output_zip = io.BytesIO()
            
            with zipfile.ZipFile(input_zip, 'r') as zin:
                with zipfile.ZipFile(output_zip, 'w') as zout:
                    for item in zin.infolist():
                        content = zin.read(item.filename)
                        
                        if item.filename == 'word/document.xml':
                            content = self._process_docx_xml_content(content, operation)
                        elif item.filename.startswith('word/header') and item.filename.endswith('.xml'):
                            content = self._process_docx_xml_content(content, operation)
                        elif item.filename.startswith('word/footer') and item.filename.endswith('.xml'):
                            content = self._process_docx_xml_content(content, operation)
                        
                        zout.writestr(item, content)
            
            return output_zip.getvalue()
        
        except Exception as e:
            self._log(f"Error in direct DOCX processing: {e}")
            raise
    
    def _process_docx_xml_content(self, xml_content, operation):
        """
        Process DOCX XML
        """
        try:
            root = ET.fromstring(xml_content)
            
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            }
            
            text_elements = root.findall('.//w:t', namespaces)
            text_elements = [elem for elem in text_elements if elem.text]
            texts = [elem.text or "" for elem in text_elements]
            processed_texts = self._process_text_batch(texts, operation)
            
            for elem, original_text, processed_text in zip(text_elements, texts, processed_texts):
                if processed_text != original_text:
                    elem.text = processed_text
            
            # Return the processed XML content
            return ET.tostring(root, encoding='utf-8', method='xml')
        
        except Exception as e:
            self._log(f"Error processing DOCX XML content: {e}")
            return xml_content
    
    def _process_docx_with_python_docx(self, file_content, operation):
        """
        Using Python docx to process DOCX files
        """
        try:
            doc = Document(io.BytesIO(file_content))
            
            paragraphs = list(doc.paragraphs)

            def _collect_table_paragraphs(table_obj, collection):
                for row in table_obj.rows:
                    for cell in row.cells:
                        collection.extend(cell.paragraphs)

            for table in doc.tables:
                _collect_table_paragraphs(table, paragraphs)

            for section in doc.sections:
                paragraphs.extend(section.header.paragraphs)
                for table in section.header.tables:
                    _collect_table_paragraphs(table, paragraphs)
                paragraphs.extend(section.footer.paragraphs)
                for table in section.footer.tables:
                    _collect_table_paragraphs(table, paragraphs)

            original_texts = [para.text or "" for para in paragraphs]
            processed_texts = self._process_text_batch(original_texts, operation)

            for para, original_text, processed_text in zip(paragraphs, original_texts, processed_texts):
                if processed_text != original_text:
                    self._apply_text_to_paragraph(para, processed_text)

            output = io.BytesIO()
            doc.save(output)
            return output.getvalue()
            
        except Exception as e:
            self._log(f"Error processing DOCX with python-docx: {e}")
            return file_content
    
    def _create_pdf_with_layout(self, content, filename, original_file_content, metadata):
        """
        Create a PDF with preserved layout
        """
        start_time = perf_counter()
        try:
            pdf_buffer = io.BytesIO()
            
            doc = SimpleDocTemplate(
                pdf_buffer,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            styles = getSampleStyleSheet()
            normal_style = styles['Normal']
            
            story = self._build_pdf_story(content, filename, styles)
            
            tables = metadata.get('tables', [])
            if tables:
                story.append(Spacer(1, 12))
                story.append(Paragraph("Extracted Tables:", styles['Heading2']))
                story.append(Spacer(1, 6))
                
                for i, table_data in enumerate(tables):
                    # if i > 2:  # 只显示前3个表格，避免PDF过大
                    #     story.append(Paragraph(f"... and {len(tables) - 3} more tables", normal_style))
                    #     break
                    
                    story.append(Paragraph(f"Table {i+1}:", styles['Heading3']))
                    
                    table_content = table_data.get('data', [])
                    if table_content:
                        max_rows = min(10, len(table_content))
                        max_cols = min(6, len(table_content[0]) if table_content else 0)
                        
                        display_data = []
                        for row_idx in range(max_rows):
                            if row_idx < len(table_content):
                                row = table_content[row_idx]
                                display_row = []
                                for col_idx in range(max_cols):
                                    if col_idx < len(row):
                                        cell_text = str(row[col_idx])[:50]  
                                        display_row.append(cell_text)
                                    else:
                                        display_row.append("")
                                display_data.append(display_row)
                        
                        if display_data:
                            table = Table(display_data)
                            table.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                ('FONTSIZE', (0, 0), (-1, 0), 10),
                                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                                ('FONTSIZE', (0, 1), (-1, -1), 8),
                                ('GRID', (0, 0), (-1, -1), 1, colors.black)
                            ]))
                            story.append(table)
                            story.append(Spacer(1, 12))
            
            images = metadata.get('images', [])
            
            if images:
                story.append(Spacer(1, 12))
                story.append(Paragraph("Extracted Images:", styles['Heading2']))
                story.append(Spacer(1, 6))
                
                image_count = 0
                for img_data in images:
                    image_bytes = img_data.get('image_data')
                    if not image_bytes:
                        continue
                    extracted_text = (img_data.get('extracted_text') or '').strip().lower()
                    if extracted_text in ["stc", "sic"]:
                        continue

                    try:
                        # Create an Image using the BytesIO object
                        img_stream = io.BytesIO(image_bytes)
                        
                        try:
                            pil_img = PILImage.open(img_stream)
                            original_width, original_height = pil_img.size
                            
                            max_width = 400
                            if original_width > max_width:
                                ratio = max_width / original_width
                                new_height = int(original_height * ratio)
                                width = max_width
                                height = new_height
                            else:
                                width = original_width
                                height = original_height
                                
                            img_stream.seek(0)
                        except Exception as img_size_error:
                            self._log(f"Could not get image dimensions, using default: {img_size_error}")
                            width = 400
                            height = 300
                            img_stream.seek(0)
                        
                        img = Image(img_stream, width=width, height=height)
                        
                        description = img_data.get('description', '')
                        if description:
                            desc_para = Paragraph(f"Image {image_count + 1}: {description}", normal_style)
                            story.append(desc_para)
                            story.append(Spacer(1, 3))
                        
                        story.append(img)
                        story.append(Spacer(1, 6))
                        
                        story.append(Spacer(1, 12))
                        image_count += 1
                    except Exception as img_error:
                        self._log(f"Error adding image to PDF: {img_error}")
                        continue
            
            # Build PDF
            doc.build(story)
            
            pdf_content = pdf_buffer.getvalue()
            pdf_buffer.close()
            
            return pdf_content
            
        except Exception as e:
            self._log(f"Error creating PDF with layout: {e}")
            # If the creation fails, go back to simple PDF creation
            return self._create_pdf(content, filename)
        finally:
            self._record_timing('pdf_generation_layout', perf_counter() - start_time)
    
    def _create_pdf(self, content, filename):
        """
        Create simple PDF file
        """
        start_time = perf_counter()
        try:
            pdf_buffer = io.BytesIO()
            
            doc = SimpleDocTemplate(
                pdf_buffer,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            styles = getSampleStyleSheet()
            story = self._build_pdf_story(content, filename, styles)
            
            # Build PDF
            doc.build(story)
            
            pdf_content = pdf_buffer.getvalue()
            pdf_buffer.close()
            
            return pdf_content
            
        except Exception as e:
            self._log(f"Error creating PDF: {e}")
            return self._create_error_pdf(f"Error creating PDF: {str(e)}")
        finally:
            self._record_timing('pdf_generation', perf_counter() - start_time)

    def _build_pdf_story(self, content, filename, styles):
        if isinstance(content, bytes):
            content = content.decode('utf-8', errors='ignore')

        title_style = styles['Heading1']
        normal_style = styles['Normal']

        story = []

        title = Paragraph(f"Processed Document: {filename}", title_style)
        story.append(title)
        story.append(Spacer(1, 12))

        time_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        time_para = Paragraph(f"Processed on: {time_str}", normal_style)
        story.append(time_para)
        story.append(Spacer(1, 12))

        story.append(Spacer(1, 6))

        paragraphs = content.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                safe_text = html.escape(para_text)
                safe_text = safe_text.replace('\n', '<br/>')

                try:
                    para = Paragraph(safe_text, normal_style)
                    story.append(para)
                    story.append(Spacer(1, 6))
                except Exception as para_error:
                    self._log(f"Error creating paragraph: {para_error}")
                    safe_text = "【The content contains characters that cannot be processed, skipped】"
                    para = Paragraph(safe_text, normal_style)
                    story.append(para)
                    story.append(Spacer(1, 6))

        return story
    
    def _create_error_pdf(self, error_message):
        """
        Create a PDF containing error information
        """
        try:
            pdf_buffer = io.BytesIO()
            doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            
            story = []
            title = Paragraph("PDF Generation Error", styles['Heading1'])
            story.append(title)
            story.append(Spacer(1, 12))
            
            error_para = Paragraph(f"Error: {error_message}", styles['Normal'])
            story.append(error_para)
            
            doc.build(story)
            pdf_content = pdf_buffer.getvalue()
            pdf_buffer.close()
            
            return pdf_content
        except:
            return f"PDF generation failed: {error_message}".encode('utf-8')
    
    def _extract_images_from_docx(self, doc, file_content):
        """
        Extract image information and data from DOCX
        """
        images_data = []
        
        try:
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.target_ref:
                    try:
                        image_part = rel.target_part
                        image_bytes = image_part.blob
                        image_info = {
                            "type": "docx_embedded_image",
                            "description": f"Embedded image {rel_id}",
                            "image_data": image_bytes,
                            "image_format": self._get_image_format(image_bytes),
                            "extracted_text": "",
                            "ocr_applied": False,
                        }

                        images_data.append(image_info)
                    except Exception as e:
                        self._log(f"Error extracting image {rel_id}: {e}")
                        images_data.append({
                            "type": "docx_embedded_image",
                            "description": f"Embedded image {rel_id} (extraction failed)",
                            "extracted_text": f"[Image extraction failed: {str(e)}]",
                            "ocr_applied": False,
                            "image_data": b"",
                        })
        except Exception as e:
            self._log(f"Error processing DOCX images: {e}")
        
        return self._process_images_with_ocr(images_data)
    
    def _get_image_format(self, image_data):
        """
        Detect image format
        """
        try:
            image = PILImage.open(io.BytesIO(image_data))
            return image.format
        except:
            return "unknown"

    def _should_apply_ocr(self, image_bytes):
        """Simple heuristic to skip tiny or low-contrast images before OCR."""
        try:
            with PILImage.open(io.BytesIO(image_bytes)) as img:
                width, height = img.size
                if width < 32 or height < 32 or width * height < 2000:
                    return False
                gray = img.convert("L")
                stats = ImageStat.Stat(gray)
                if stats.stddev[0] < 8:
                    return False
                sample = gray.resize((64, 64))
                pixels = sample.getdata()
                dark_pixels = sum(1 for p in pixels if p < 110)
                if dark_pixels / len(pixels) < 0.01:
                    return False
                return True
        except Exception as exc:
            self._log(f"OCR heuristic fallback to processing: {exc}")
            return True

    def _prepare_image_array_for_ocr(self, image_bytes):
        """Convert bytes to a grayscale numpy array for OCR."""
        with PILImage.open(io.BytesIO(image_bytes)) as image:
            gray = image.convert("L")
            return np.array(gray)

    def _process_images_with_ocr(self, images):
        """Apply OCR to images with simple parallelism and ordering preserved."""
        if not images:
            return []

        if self.throughput_mode or not self.ocr_enabled:
            reason = "[OCR disabled in throughput mode]" if self.throughput_mode else "[OCR disabled]"
            for img in images:
                img["extracted_text"] = reason
                img["ocr_applied"] = False
            self._ocr_images_skipped += len(images)
            return images

        if not self.ocr_available or not self._paddle_ocr:
            for img in images:
                img["extracted_text"] = "[OCR not available]"
                img["ocr_applied"] = False
            self._ocr_images_skipped += len(images)
            return images

        workers = min(2, os.cpu_count() or 2)
        futures = {}

        with ThreadPoolExecutor(max_workers=workers) as executor:
            for idx, img in enumerate(images):
                if self._ocr_images_processed + len(futures) >= OCR_MAX_IMAGES_PER_DOC:
                    img["extracted_text"] = "[OCR skipped: max images reached]"
                    img["ocr_applied"] = False
                    self._ocr_images_skipped += 1
                    continue
                if not img.get("image_data"):
                    img["extracted_text"] = "[OCR skipped: no image data]"
                    img["ocr_applied"] = False
                    self._ocr_images_skipped += 1
                    continue
                if not self._should_apply_ocr(img.get("image_data", b"")):
                    img["extracted_text"] = "[OCR skipped: low-text likelihood]"
                    img["ocr_applied"] = False
                    self._ocr_images_skipped += 1
                    continue

                futures[executor.submit(self._perform_ocr, img["image_data"])] = idx

            for future in as_completed(futures):
                idx = futures[future]
                try:
                    extracted = future.result()
                    images[idx]["extracted_text"] = extracted
                    images[idx]["ocr_applied"] = True
                    self._ocr_images_processed += 1
                except Exception as exc:
                    images[idx]["extracted_text"] = f"[OCR failed: {exc}]"
                    images[idx]["ocr_applied"] = False
                    self._ocr_images_skipped += 1

        return images

    def _perform_ocr(self, image_data):
        """
        Perform OCR processing using PaddleOCR.
        """
        start_time = perf_counter()
        try:
            if not self._paddle_ocr:
                return "[OCR not available]"

            image_array = self._prepare_image_array_for_ocr(image_data)
            ocr_result = self._paddle_ocr.ocr(image_array, cls=False)
            texts = []
            for line in ocr_result or []:
                for entry in line:
                    if entry and len(entry) > 1 and entry[1]:
                        texts.append(str(entry[1][0]))
            cleaned_text = self._clean_ocr_text(" ".join(texts))
            return cleaned_text if cleaned_text else "[No text detected in image]"
        except Exception as e:
            self._log(f"OCR processing failed: {e}")
            return f"[OCR failed: {str(e)}]"
        finally:
            self._record_timing('ocr', perf_counter() - start_time)

    def _clean_ocr_text(self, text):
        """
        Clean up OCR extracted text
        """
        text = text or ""
        # Remove unnecessary spaces and line breaks
        text = re.sub(r'\s+', ' ', text)
        
        # Remove blank spaces at the beginning and end
        text = text.strip()
        
        return text
    
    def _extract_to_json(self, content, filename, file_extension, processing_info, metadata, original_file_content=None):
        """Extract content to JSON format"""
        json_output = deepcopy(JSON_SCHEMA)
        
        if isinstance(content, bytes):
            content = content.decode('utf-8', errors='ignore')
        
        if isinstance(original_file_content, (bytes, bytearray)):
            file_size_bytes = len(original_file_content)
        else:
            file_size_bytes = len(str(original_file_content or "").encode('utf-8', errors='ignore'))

        json_output["document_metadata"]["filename"] = filename
        json_output["document_metadata"]["file_type"] = file_extension
        json_output["document_metadata"]["processing_date"] = datetime.now().isoformat()
        json_output["document_metadata"]["file_size"] = file_size_bytes
        
        json_output["content"]["text"] = content
        json_output["content"]["tables"] = metadata.get('tables', [])
        
        json_output["content"]["images"] = self._extract_image_info(metadata.get('images', []), content)
        
        json_output["processing_info"] = processing_info
        
        return json.dumps(json_output, indent=2, ensure_ascii=False)
    
    def _extract_image_info(self, images_metadata, document_content):
        """
        Extract image information
        """
        images_info = []
        
        for img_meta in images_metadata:
            image_info = {
                "type": img_meta.get("type", "unknown"),
                "description": img_meta.get("description", ""),
                "extracted_text": img_meta.get("extracted_text", ""),
                "ocr_applied": img_meta.get("ocr_applied", False),
                "image_format": img_meta.get("image_format", "unknown")
            }
            
            # If the image data is large, not to include it in JSON
            # or, include Base64 encoded thumbnails
            if "image_data" in img_meta and len(img_meta["image_data"]) < 10000:  # Only includes small images
                try:
                    image = PILImage.open(io.BytesIO(img_meta["image_data"]))
                    # Create Thumbnail 
                    image.thumbnail((100, 100))
                    thumb_buffer = io.BytesIO()
                    image.save(thumb_buffer, format="PNG")
                    image_info["thumbnail"] = base64.b64encode(thumb_buffer.getvalue()).decode('utf-8')
                except:
                    pass  # Ignore thumbnail creation error
            
            images_info.append(image_info)
        
        # Identify image references from document text
        if isinstance(document_content, bytes):
            document_content = document_content.decode('utf-8', errors='ignore')
            
        image_refs = re.findall(r'\[Image:\s*(.*?)\]|\!\[(.*?)\]', document_content)
        for match in image_refs:
            alt_text = match[0] or match[1]
            if alt_text and alt_text not in [img.get("description", "") for img in images_info]:
                images_info.append({
                    "type": "referenced_image",
                    "description": alt_text,
                    "extracted_text": f"Referenced image: {alt_text}",
                    "ocr_applied": False
                })
        
        # remove stc logo
        rtn_images = []
        for image in images_info:
            text_value = (image.get("extracted_text") or "").lower()
            if text_value in ["stc", "sic"]:
                continue
            rtn_images.append(image)

        return rtn_images

    def _extract_image_with_pymupdf(self, doc, page_num, img_idx, bbox=None, xref=None):
        """Extract an image region or direct xref using PyMuPDF."""
        try:
            page = doc[page_num]

            if xref is not None:
                pix = fitz.Pixmap(doc, xref)
                try:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                except Exception:
                    if pix.n - pix.alpha > 3:  # Change CMYK to RGB as fallback
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                return pix

            if bbox is not None:
                try:
                    rect = fitz.Rect(bbox)
                    pix = page.get_pixmap(
                        matrix=fitz.Matrix(OCR_RENDER_SCALE, OCR_RENDER_SCALE),
                        clip=rect,
                        colorspace=fitz.csGRAY,
                    )
                    return pix
                except Exception:
                    pass

            image_list = page.get_images()
            if img_idx < len(image_list):
                xref = image_list[img_idx][0]
                pix = fitz.Pixmap(doc, xref)
                if pix.n - pix.alpha > 3:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                return pix

        except Exception as e:
            self._log(f"Error extracting with PyMuPDF: {e}")

        return None
